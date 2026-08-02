from __future__ import annotations

import json
from copy import deepcopy
from pathlib import Path
from typing import Any

from docx import Document
from openpyxl import Workbook
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfbase.ttfonts import TTFError, TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


class ReportService:
    def __init__(self, report_dir: Path) -> None:
        self.report_dir = report_dir

    def save_all_reports(self, review_id: str, report: dict[str, Any]) -> dict[str, Path]:
        self.report_dir.mkdir(parents=True, exist_ok=True)
        report = self._snapshot_report(review_id, report)
        return {
            "json": self.save_json_report(review_id, report),
            "docx": self.save_docx_report(review_id, report),
            "pdf": self.save_pdf_report(review_id, report),
            "markdown": self.save_markdown_report(review_id, report),
            "xlsx": self.save_excel_report(review_id, report),
        }

    def save_json_report(self, review_id: str, report: dict[str, Any]) -> Path:
        report_path = self.report_dir / f"{review_id}.json"
        report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
        return report_path

    def save_markdown_report(self, review_id: str, report: dict[str, Any]) -> Path:
        path = self.report_dir / f"{review_id}.md"
        score = report.get("风险评分", {})
        lines = [
            "# 合同智能审查报告",
            "",
            f"- 审查编号：{report.get('审查编号', review_id)}",
            f"- 文件名：{report.get('文件名', '')}",
            f"- 总体风险等级：{report.get('总体风险等级', '')}",
            f"- 风险分：{score.get('风险分', '-')}",
            "",
            "## 审查摘要",
            "",
            str(report.get("审查摘要", "")),
            "",
            "## 风险点",
        ]
        for item in report.get("风险点", []):
            legal_text = self._legal_basis_text(item)
            lines.extend(
                [
                    "",
                    f"### {item.get('风险编号', '')} {item.get('风险标题', '')}",
                    f"- 风险等级：{item.get('风险等级', '')}",
                    f"- 风险条款：{item.get('相关条款', '')}",
                    f"- 问题说明：{item.get('问题说明', '')}",
                    f"- 可能后果：{item.get('可能后果', '')}",
                    f"- 法律依据：{legal_text}",
                    f"- 修改方向：{item.get('修改方向', '')}",
                    f"- 人工复核：{item.get('人工复核状态', '待确认')}",
                ]
            )
        lines.extend(["", "## 免责声明", "", str(report.get("免责声明", ""))])
        path.write_text("\n".join(lines), encoding="utf-8")
        return path

    def save_excel_report(self, review_id: str, report: dict[str, Any]) -> Path:
        path = self.report_dir / f"{review_id}.xlsx"
        workbook = Workbook()
        summary = workbook.active
        summary.title = "审查概览"
        score = report.get("风险评分", {})
        summary.append(["审查编号", report.get("审查编号", review_id)])
        summary.append(["文件名", report.get("文件名", "")])
        summary.append(["总体风险等级", report.get("总体风险等级", "")])
        summary.append(["风险分", score.get("风险分", "")])
        summary.append(["安全分", score.get("安全分", "")])
        risks = workbook.create_sheet("风险明细")
        risks.append(["风险编号", "等级", "类别", "标题", "相关条款", "问题说明", "可能后果", "法律依据", "修改方向", "人工复核状态"])
        for item in report.get("风险点", []):
            risks.append(
                [
                    item.get("风险编号", ""),
                    item.get("风险等级", ""),
                    item.get("风险类别", ""),
                    item.get("风险标题", ""),
                    item.get("相关条款", ""),
                    item.get("问题说明", ""),
                    item.get("可能后果", ""),
                    self._legal_basis_text(item),
                    item.get("修改方向", ""),
                    item.get("人工复核状态", "待确认"),
                ]
            )
        suggestions = workbook.create_sheet("修改建议")
        suggestions.append(["对应风险编号", "风险类别", "修改建议", "建议条款"])
        for item in report.get("修改建议", []):
            suggestions.append(
                [
                    item.get("对应风险编号", ""),
                    item.get("风险类别", ""),
                    item.get("修改建议", ""),
                    item.get("建议条款", ""),
                ]
            )
        for sheet in workbook.worksheets:
            sheet.freeze_panes = "A2"
            sheet.auto_filter.ref = sheet.dimensions
            for column in sheet.columns:
                letter = column[0].column_letter
                sheet.column_dimensions[letter].width = min(
                    60, max(12, max(len(str(cell.value or "")) for cell in column) + 2)
                )
        workbook.save(str(path))
        return path

    def save_docx_report(self, review_id: str, report: dict[str, Any]) -> Path:
        path = self.report_dir / f"{review_id}.docx"
        doc = Document()
        doc.add_heading("合同智能审查报告", level=1)
        doc.add_paragraph(f"审查编号：{report.get('审查编号', review_id)}")
        doc.add_paragraph(f"文件名：{report.get('文件名', '')}")
        doc.add_paragraph(f"总体风险等级：{report.get('总体风险等级', '')}")
        score = report.get("风险评分", {})
        doc.add_paragraph(f"风险分：{score.get('风险分', '-')}")
        doc.add_paragraph(f"安全分：{score.get('安全分', '-')}")
        doc.add_heading("审查摘要", level=2)
        doc.add_paragraph(str(report.get("审查摘要", "")))
        doc.add_heading("风险点", level=2)
        for item in report.get("风险点", []):
            doc.add_paragraph(
                f"{item.get('风险编号', '')} {item.get('风险等级', '')} {item.get('风险标题', '')}",
                style="List Bullet",
            )
            doc.add_paragraph(str(item.get("问题说明", "")))
            doc.add_paragraph(f"风险条款：{item.get('相关条款', '')}")
            doc.add_paragraph(f"可能后果：{item.get('可能后果', '')}")
            doc.add_paragraph(f"法律依据：{self._legal_basis_text(item)}")
            doc.add_paragraph(f"人工复核：{item.get('人工复核状态', '待确认')}")
        doc.add_heading("修改建议", level=2)
        for item in report.get("修改建议", []):
            doc.add_paragraph(str(item.get("修改建议", "")), style="List Bullet")
            if item.get("建议条款"):
                doc.add_paragraph(str(item.get("建议条款")))
        doc.add_heading("依据检索", level=2)
        for item in report.get("依据检索", []):
            doc.add_paragraph(
                f"{item.get('来源', '')}：{item.get('内容', '')}", style="List Bullet"
            )
        doc.add_heading("免责声明", level=2)
        doc.add_paragraph(str(report.get("免责声明", "")))
        doc.save(str(path))
        return path

    def save_pdf_report(self, review_id: str, report: dict[str, Any]) -> Path:
        path = self.report_dir / f"{review_id}.pdf"
        font_name = self._register_chinese_font()
        styles = getSampleStyleSheet()
        normal = ParagraphStyle(
            "ChineseNormal", parent=styles["Normal"], fontName=font_name, fontSize=9, leading=15
        )
        title = ParagraphStyle(
            "ChineseTitle", parent=styles["Title"], fontName=font_name, fontSize=18, leading=24
        )
        heading = ParagraphStyle(
            "ChineseHeading",
            parent=styles["Heading2"],
            fontName=font_name,
            fontSize=13,
            leading=18,
        )
        story: list[Any] = [Paragraph("合同智能审查报告", title), Spacer(1, 8)]
        score = report.get("风险评分", {})
        summary_rows = [
            ["审查编号", str(report.get("审查编号", review_id))],
            ["文件名", str(report.get("文件名", ""))],
            ["总体风险", str(report.get("总体风险等级", ""))],
            ["风险分", str(score.get("风险分", "-"))],
            ["安全分", str(score.get("安全分", "-"))],
        ]
        story.append(self._table(summary_rows, font_name))
        story.extend(
            [
                Spacer(1, 10),
                Paragraph("审查摘要", heading),
                Paragraph(str(report.get("审查摘要", "")), normal),
            ]
        )
        story.extend([Spacer(1, 10), Paragraph("风险点", heading)])
        for item in report.get("风险点", [])[:20]:
            story.append(
                Paragraph(
                    f"{item.get('风险编号', '')} [{item.get('风险等级', '')}] {item.get('风险标题', '')}",
                    normal,
                )
            )
            story.append(Paragraph(str(item.get("问题说明", "")), normal))
            story.append(Paragraph(f"风险条款：{item.get('相关条款', '')}", normal))
            story.append(Paragraph(f"可能后果：{item.get('可能后果', '')}", normal))
            story.append(Paragraph(f"法律依据：{self._legal_basis_text(item)}", normal))
            story.append(Paragraph(f"人工复核：{item.get('人工复核状态', '待确认')}", normal))
            story.append(Spacer(1, 5))
        story.extend([Spacer(1, 10), Paragraph("修改建议", heading)])
        for item in report.get("修改建议", [])[:20]:
            story.append(Paragraph(str(item.get("修改建议", "")), normal))
            if item.get("建议条款"):
                story.append(Paragraph(str(item.get("建议条款")), normal))
            story.append(Spacer(1, 5))
        story.extend(
            [
                Spacer(1, 10),
                Paragraph("免责声明", heading),
                Paragraph(str(report.get("免责声明", "")), normal),
            ]
        )
        document = SimpleDocTemplate(
            str(path),
            pagesize=A4,
            rightMargin=16 * mm,
            leftMargin=16 * mm,
            topMargin=16 * mm,
            bottomMargin=16 * mm,
        )
        document.build(story)
        return path

    def _snapshot_report(self, review_id: str, report: dict[str, Any]) -> dict[str, Any]:
        """Freeze human-readable citation fields so later knowledge changes cannot rewrite history."""
        snapshot = deepcopy(report)
        snapshot["审查编号"] = review_id
        snapshot.setdefault(
            "免责声明",
            "本报告由人工智能与规则引擎辅助生成，仅供合同风险初筛和内部决策参考，不构成正式法律意见；重要合同请由专业法律人员复核。",
        )
        for item in snapshot.get("风险点", []):
            if not isinstance(item, dict):
                continue
            raw_basis = item.get("legalBasis") or item.get("legal_basis") or []
            verified = []
            for basis in raw_basis:
                if not isinstance(basis, dict):
                    continue
                verified.append(
                    {
                        "legalArticleId": basis.get("legalArticleId"),
                        "法律名称快照": basis.get("lawName", ""),
                        "条号快照": basis.get("articleNo", ""),
                        "引用文本快照": basis.get("contentSummary", ""),
                        "来源地址快照": basis.get("sourceUrl", ""),
                    }
                )
            item["法律依据快照"] = verified
            item.setdefault("人工复核状态", "待确认")
        return snapshot

    def _legal_basis_text(self, item: dict[str, Any]) -> str:
        basis_items = item.get("法律依据快照") or item.get("legalBasis") or item.get("legal_basis") or []
        values = []
        for basis in basis_items:
            if not isinstance(basis, dict):
                continue
            law_name = basis.get("法律名称快照") or basis.get("lawName") or ""
            article_no = basis.get("条号快照") or basis.get("articleNo") or ""
            if law_name or article_no:
                values.append(f"《{law_name}》{article_no}")
        return "；".join(values) or "未匹配到已核验法律依据"

    def _table(self, rows: list[list[str]], font_name: str) -> Table:
        table = Table(rows, colWidths=[32 * mm, 130 * mm])
        table.setStyle(
            TableStyle(
                [
                    ("FONTNAME", (0, 0), (-1, -1), font_name),
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#9aa4b2")),
                    ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#eef2f7")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ]
            )
        )
        return table

    def _register_chinese_font(self) -> str:
        if "ChineseFont" in pdfmetrics.getRegisteredFontNames():
            return "ChineseFont"
        candidates = [
            Path("C:/Windows/Fonts/msyh.ttc"),
            Path("C:/Windows/Fonts/simhei.ttf"),
            Path("C:/Windows/Fonts/simsun.ttc"),
            Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"),
            Path("/usr/share/fonts/opentype/noto/NotoSansCJKsc-Regular.otf"),
        ]
        for candidate in candidates:
            if candidate.exists():
                try:
                    pdfmetrics.registerFont(TTFont("ChineseFont", str(candidate)))
                except (OSError, TTFError):
                    continue
                return "ChineseFont"

        fallback = "STSong-Light"
        if fallback not in pdfmetrics.getRegisteredFontNames():
            pdfmetrics.registerFont(UnicodeCIDFont(fallback))
        return fallback

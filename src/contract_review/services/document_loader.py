from __future__ import annotations

import shutil
import subprocess
import tempfile
from pathlib import Path

from PIL import Image

from contract_review.core.config import Settings
from contract_review.core.exceptions import (
    DocumentTextExtractionError,
    UnsupportedDocumentTypeError,
)

NO_EXTRACTABLE_TEXT_MESSAGE = (
    "当前文件可能为扫描版或无可提取文本，请上传可复制文字的 PDF 或 Word 文件。"
)


class DocumentLoader:
    supported_extensions = {
        ".pdf",
        ".doc",
        ".docx",
        ".txt",
        ".png",
        ".jpg",
        ".jpeg",
        ".tif",
        ".tiff",
        ".bmp",
    }

    def __init__(self, settings: Settings) -> None:
        self.settings = settings

    def load_text(self, file_path: Path) -> str:
        suffix = file_path.suffix.lower()
        if suffix not in self.supported_extensions:
            raise UnsupportedDocumentTypeError(f"Unsupported contract file type: {suffix}")
        if not file_path.is_file():
            raise UnsupportedDocumentTypeError(
                f"Unsupported contract file type or missing upload: {suffix}"
            )

        if suffix == ".pdf":
            text = self._load_pdf(file_path)
        elif suffix == ".docx":
            text = self._load_docx(file_path)
        elif suffix == ".doc":
            text = self._load_legacy_doc(file_path)
        elif suffix == ".txt":
            text = self._load_txt(file_path)
        else:
            text = self._load_image_with_ocr(file_path)
        if len(text.strip()) < 2:
            raise DocumentTextExtractionError(NO_EXTRACTABLE_TEXT_MESSAGE)
        return text.strip()

    def _load_pdf(self, file_path: Path) -> str:
        import fitz

        with fitz.open(file_path) as document:
            pages = [page.get_text("text").strip() for page in document]
            if sum(len(page) for page in pages) >= max(80, len(document) * 20):
                return "\n".join(pages)
            ocr_pages = []
            for page in document:
                pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                image = Image.frombytes("RGB", (pixmap.width, pixmap.height), pixmap.samples)
                ocr_pages.append(self._ocr_image(image))
            return "\n".join(ocr_pages)

    def _load_docx(self, file_path: Path) -> str:
        from docx import Document

        document = Document(file_path)
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        table_text = []
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    table_text.append(" | ".join(cells))
        return "\n".join([*paragraphs, *table_text])

    def _load_txt(self, file_path: Path) -> str:
        raw = file_path.read_bytes()
        for encoding in ("utf-8-sig", "utf-16", "gb18030"):
            try:
                return raw.decode(encoding)
            except UnicodeDecodeError:
                continue
        raise DocumentTextExtractionError(NO_EXTRACTABLE_TEXT_MESSAGE)

    def _load_image_with_ocr(self, file_path: Path) -> str:
        with Image.open(file_path) as image:
            return self._ocr_image(image)

    def _ocr_image(self, image: Image.Image) -> str:
        import pytesseract

        if self.settings.tesseract_cmd:
            pytesseract.pytesseract.tesseract_cmd = self.settings.tesseract_cmd
        config = (
            f'--tessdata-dir "{self.settings.tessdata_dir}"' if self.settings.tessdata_dir else ""
        )
        return pytesseract.image_to_string(image, lang=self.settings.ocr_languages, config=config)

    def _load_legacy_doc(self, file_path: Path) -> str:
        command = self.settings.libreoffice_cmd or shutil.which("soffice")
        if not command:
            raise UnsupportedDocumentTypeError(
                "读取旧版 .doc 需要安装 LibreOffice，并配置 LIBREOFFICE_CMD"
            )
        with tempfile.TemporaryDirectory() as output_dir:
            completed = subprocess.run(
                [
                    command,
                    "--headless",
                    "--convert-to",
                    "docx",
                    "--outdir",
                    output_dir,
                    str(file_path),
                ],
                capture_output=True,
                text=True,
                timeout=60,
                check=False,
            )
            converted = Path(output_dir) / f"{file_path.stem}.docx"
            if completed.returncode != 0 or not converted.exists():
                raise UnsupportedDocumentTypeError("旧版 .doc 转换失败，请检查 LibreOffice 配置")
            return self._load_docx(converted)

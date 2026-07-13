from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient

from contract_review.core.config import get_settings
from contract_review.main import create_app


def _auth_headers(client: TestClient) -> dict[str, str]:
    response = client.post(
        "/api/v1/auth/login",
        json={"email": "admin@example.com", "password": "Admin12345!"},
    )
    return {"Authorization": f"Bearer {response.json()['data']['access_token']}"}


def _make_contract(path: Path) -> None:
    document = Document()
    document.add_heading("采购合同", level=1)
    document.add_paragraph("甲方：北京示例科技有限公司")
    document.add_paragraph("乙方：上海测试服务有限公司")
    document.add_paragraph("合同金额：人民币100000元")
    document.add_paragraph("履行期限：2026年8月1日至2026年12月31日")
    document.add_paragraph("甲方应在验收合格并收到发票后10个工作日内付款。")
    document.save(path)


def test_review_docx_flow(tmp_path: Path, monkeypatch) -> None:
    monkeypatch.setenv("ENABLE_LLM", "false")
    get_settings.cache_clear()

    contract_path = tmp_path / "采购合同.docx"
    _make_contract(contract_path)

    with TestClient(create_app()) as client:
        headers = _auth_headers(client)
        with contract_path.open("rb") as file_obj:
            response = client.post(
                "/api/v1/reviews",
                files={
                    "合同文件": (
                        contract_path.name,
                        file_obj,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
                headers=headers,
            )

    payload = response.json()
    assert response.status_code == 201
    assert payload["status"] == "已完成"
    assert "甲方：北京示例科技有限公司" in payload["contract_text"]
    assert payload["extracted_fields"]["合同主体"]
    assert payload["extracted_fields"]["合同金额"]
    assert payload["final_report"]["总体风险等级"] in {"低风险", "中风险", "高风险"}
    assert "风险评分" in payload["final_report"]
    assert payload["final_report"]["依据检索"]
    assert payload["report_path"]
    assert payload["export_paths"]["json"]
    assert payload["export_paths"]["docx"]
    assert payload["export_paths"]["pdf"]
    assert all("原文定位" in finding for finding in payload["risk_findings"])


def test_review_history_and_download(tmp_path: Path, monkeypatch) -> None:
    monkeypatch.setenv("ENABLE_LLM", "false")
    get_settings.cache_clear()
    contract_path = tmp_path / "采购合同.docx"
    _make_contract(contract_path)

    with TestClient(create_app()) as client:
        headers = _auth_headers(client)
        with contract_path.open("rb") as file_obj:
            create_response = client.post(
                "/api/v1/reviews",
                files={
                    "合同文件": (
                        contract_path.name,
                        file_obj,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
                headers=headers,
            )
        review_id = create_response.json()["review_id"]
        history_response = client.get("/api/v1/reviews", headers=headers)
        report_response = client.get(f"/api/v1/reviews/{review_id}", headers=headers)
        download_response = client.get(
            f"/api/v1/reviews/{review_id}/download?file_type=json", headers=headers
        )

    assert history_response.status_code == 200
    assert any(item["review_id"] == review_id for item in history_response.json())
    assert report_response.status_code == 200
    assert download_response.status_code == 200


def test_review_upload_and_download_require_authentication() -> None:
    with TestClient(create_app()) as client:
        assert client.get("/api/v1/reviews").status_code == 401
        assert client.get("/api/v1/reviews/unknown/download?file_type=json").status_code == 401
